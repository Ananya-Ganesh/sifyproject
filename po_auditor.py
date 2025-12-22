"""
AI Purchase Order Auditor - Multi-format ENC (Extract-Normalize-Compare) System

- Handles PDF (digital & scanned), Images (PNG/JPG), Word (.docx), Excel (.xlsx)
- Normalizes into unified schema:
    {
      "line_items": [{"name": str, "quantity": int, "unit_price": float, "total": float}],
      "otc": float, "arc": float, "grand_total": float
    }
- Applies 6-point audit logic with 70% product overlap gate.
"""

from __future__ import annotations

import os
import io
import re
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional, Any

import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document as DocxDocument
import pandas as pd

from rapidfuzz import fuzz, process
from word2number import w2n
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ---------------------- Unified Schema ---------------------- #

@dataclass
class LineItem:
    name: str
    quantity: int
    unit_price: float
    total: float


@dataclass
class PurchaseOrder:
    order_id: str
    line_items: List[LineItem]
    otc: float
    arc: float
    grand_total: float


# ---------------------- Normalizers ---------------------- #

class NumericNormalizer:
    """Handles digits, number words, and OCR artefacts."""

    @staticmethod
    def _clean_ocr_digits(text: str) -> str:
        # Common OCR fixes: O->0, l->1, I->1 in number-like regions
        def fix_token(tok: str) -> str:
            if re.search(r"\d", tok) or re.fullmatch(r"[OIl]+", tok):
                tok = tok.replace("O", "0").replace("o", "0")
                tok = tok.replace("l", "1").replace("I", "1")
            return tok

        return " ".join(fix_token(t) for t in text.split())

    def normalize_number(self, value: Any) -> float:
        if value is None:
            return 0.0
        if isinstance(value, (int, float)):
            return float(value)

        text = str(value).strip()
        if not text:
            return 0.0

        text = self._clean_ocr_digits(text)

        # Remove currency symbols, commas
        digits = re.sub(r"[^\d.\-]", "", text.replace(",", ""))
        if digits:
            try:
                return float(digits)
            except ValueError:
                pass

        # Try number words
        try:
            n = w2n.word_to_num(text.lower())
            return float(n)
        except Exception:
            return 0.0


class TextNormalizer:
    """Handles semantic normalization of product names, dates, etc."""

    @staticmethod
    def normalize_name(name: str) -> str:
        if not name:
            return ""
        name = name.lower()
        # Preserve numeric product codes (do not strip leading digits)
        # Replace separators with spaces
        name = re.sub(r"[_/]", " ", name)
        # Basic synonyms/shortforms
        synonyms = {
            "hb pencil": "pencil hb",
            "pencil hb": "pencil hb",
            "long size scale": "long scale",
            "longscale": "long scale",
            "ill mbps": "ill connectivity mbps",
            "ill connectivity": "ill connectivity mbps",
            "mbps service": "ill connectivity mbps",
            "ill connectivity mbps": "ill connectivity mbps",
        }
        for k, v in synonyms.items():
            if k in name:
                name = name.replace(k, v)
        # Collapse whitespace
        name = re.sub(r"\s+", " ", name).strip()
        return name

    @staticmethod
    def product_similarity(a: str, b: str) -> float:
        na = TextNormalizer.normalize_name(a)
        nb = TextNormalizer.normalize_name(b)
        if not na or not nb:
            return 0.0
        # combine fuzzy metrics
        ts = fuzz.token_set_ratio(na, nb) / 100.0
        pr = fuzz.partial_ratio(na, nb) / 100.0
        ra = fuzz.ratio(na, nb) / 100.0
        return 0.5 * ts + 0.3 * pr + 0.2 * ra


# ---------------------- Multi-format Extraction ---------------------- #

class FileExtractor:
    """Extract raw text and simple tables from various file formats."""

    def extract(self, path: str) -> Tuple[str, List[List[str]]]:
        ext = os.path.splitext(path)[1].lower()
        if ext in [".pdf"]:
            return self._extract_from_pdf(path)
        if ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff"]:
            return self._extract_from_image(path)
        if ext in [".docx"]:
            return self._extract_from_docx(path)
        if ext in [".xlsx"]:
            return self._extract_from_xlsx(path)
        # Fallback: treat as text
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            txt = f.read()
        return txt, []

    # ---- PDF ----
    def _extract_from_pdf(self, path: str) -> Tuple[str, List[List[str]]]:
        text_blocks: List[str] = []
        tables: List[List[str]] = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                # Try to see if it has extractable text
                txt = page.extract_text() or ""
                text_blocks.append(txt)

                # Capture tables as simple CSV-like rows
                page_tables = page.extract_tables() or []
                for tbl in page_tables:
                    for row in tbl:
                        tables.append([str(c) if c is not None else "" for c in row])

        full_text = "\n".join(text_blocks).strip()

        # If PDF has almost no selectable text, fall back to OCR
        if len(full_text.replace("\n", "").strip()) < 20:
            return self._ocr_pdf(path)
        return full_text, tables

    def _ocr_pdf(self, path: str) -> Tuple[str, List[List[str]]]:
        images = convert_from_path(path)
        texts = []
        for img in images:
            txt = pytesseract.image_to_string(img)
            texts.append(txt)
        return "\n".join(texts), []

    # ---- Images ----
    def _extract_from_image(self, path: str) -> Tuple[str, List[List[str]]]:
        img = Image.open(path)
        txt = pytesseract.image_to_string(img)
        return txt, []

    # ---- DOCX ----
    def _extract_from_docx(self, path: str) -> Tuple[str, List[List[str]]]:
        doc = DocxDocument(path)
        texts: List[str] = []
        tables: List[List[str]] = []

        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text)

        for tbl in doc.tables:
            for row in tbl.rows:
                row_vals = [cell.text.strip() for cell in row.cells]
                tables.append(row_vals)

        return "\n".join(texts), tables

    # ---- XLSX ----
    def _extract_from_xlsx(self, path: str) -> Tuple[str, List[List[str]]]:
        texts: List[str] = []
        tables: List[List[str]] = []

        # Use pandas.read_excel to support complex sheets and engines.
        try:
            sheets = pd.read_excel(path, sheet_name=None, engine=None)
        except Exception:
            # Last-resort: try ExcelFile (older pandas) which may sometimes work
            try:
                xls = pd.ExcelFile(path)
                sheets = {name: xls.parse(name) for name in xls.sheet_names}
            except Exception:
                return "", []

        for sheet_name, df in sheets.items():
            if df is None or df.empty:
                continue

            # Normalize MultiIndex columns and non-string headers
            try:
                cols = list(df.columns)
            except Exception:
                cols = [str(c) for c in df.columns]

            headers = [str(h) for h in cols]
            tables.append(headers)

            # Convert all cells to string safely and preserve empty cells
            for row in df.fillna("").itertuples(index=False, name=None):
                vals = [str(v) if v is not None else "" for v in row]
                tables.append(vals)
                texts.append(" ".join(vals))

        return "\n".join(texts), tables


# ---------------------- Schema Mapping (The Bridge) ---------------------- #

class DocumentExtractor:
    """
    Maps raw text/tables from any format into PurchaseOrder (Unified Schema).
    """

    def __init__(self):
        self.num_norm = NumericNormalizer()
        self.text_norm = TextNormalizer()

    # Public entry
    def extract_po(self, path: str, order_id: Optional[str] = None) -> PurchaseOrder:
        raw_text, tables = FileExtractor().extract(path)
        if not order_id:
            order_id = os.path.basename(path)

        # Try structured tables first
        line_items: List[LineItem] = self._extract_line_items_from_tables(tables)
        if not line_items:
            line_items = self._extract_line_items_from_text(raw_text)

        # Try to detect OTC, ARC, Grand Total from tables first, then text
        otc, arc, grand_total = self._extract_totals(raw_text, tables)

        # If no explicit grand total was found, but we have line items,
        # treat the sum of line item totals as the grand total. This
        # prevents false conflicts when two POs have identical line items
        # but only one explicitly prints a "Grand Total" row.
        if grand_total == 0.0 and line_items:
            grand_total = sum(li.total for li in line_items)

        return PurchaseOrder(
            order_id=order_id,
            line_items=line_items,
            otc=otc,
            arc=arc,
            grand_total=grand_total,
        )

    # ---- Tables ----
    def _extract_line_items_from_tables(self, tables: List[List[str]]) -> List[LineItem]:
        line_items: List[LineItem] = []
        if not tables:
            return line_items

        # Try to find header row with keywords
        header_idx: Optional[int] = None
        for i, row in enumerate(tables[:10]):
            joined = " ".join(row).lower()
            if any(k in joined for k in ["description", "item", "product", "service description", "desc"]) and \
               any(k in joined for k in ["qty", "quantity", "qty.", "qtty"]) and \
               any(k in joined for k in ["amount", "price", "rate", "total", "total amount", "amount (rs)", "line amount", "total (rs)", "value"]):
                header_idx = i
                break

        # If no explicit header row detected, assume first row is header
        if header_idx is None:
            header_idx = 0

        headers = [h.lower() for h in tables[header_idx]]
        rows = tables[header_idx + 1:]

        # Map columns
        def col_idx(cands: List[str]) -> Optional[int]:
            for cand in cands:
                for i, h in enumerate(headers):
                    if cand in h:
                        return i
            return None

        idx_name = col_idx(["description", "item", "product", "name", "service description", "desc"])
        idx_qty = col_idx(["qty", "quantity", "qty.", "qtty"])
        idx_unit = col_idx(["unit price", "unit_price", "price", "rate", "unit cost", "cost"])
        idx_total = col_idx(["amount", "total", "line total", "total amount", "amount (rs)", "total (rs)", "value"])

        # Keywords that indicate metadata / header rows rather than real products
        metadata_keywords = [
            "internal po reference",
            "po number",
            "customer po number",
            "linked customer po",
            "issue date",
            "po date",
            "reference",
        ]

        for r in rows:
            if idx_name is None:
                continue

            name = r[idx_name] if idx_name < len(r) else ""
            if not str(name).strip():
                continue

            qty = self.num_norm.normalize_number(r[idx_qty] if idx_qty is not None and idx_qty < len(r) else 0)
            unit_val = self.num_norm.normalize_number(
                r[idx_unit] if idx_unit is not None and idx_unit < len(r) else 0
            )
            total_val = self.num_norm.normalize_number(
                r[idx_total] if idx_total is not None and idx_total < len(r) else 0
            )

            # If numeric fields are missing/zero, try to infer from all numbers in the row
            row_numbers = re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", " ".join(r))
            if (qty == 0 or unit_val == 0 or total_val == 0) and len(row_numbers) >= 3:
                try:
                    qty = qty or self.num_norm.normalize_number(row_numbers[-3])
                    unit_val = unit_val or self.num_norm.normalize_number(row_numbers[-2])
                    total_val = total_val or self.num_norm.normalize_number(row_numbers[-1])
                except Exception:
                    pass

            # If only total is present, infer unit price if qty>0
            if unit_val == 0 and qty > 0 and total_val > 0:
                unit_val = total_val / qty

            li = LineItem(
                name=self.text_norm.normalize_name(str(name)),
                quantity=int(qty) if qty else 0,
                unit_price=unit_val,
                total=total_val if total_val else unit_val * qty,
            )

            # Skip pure metadata rows (no numeric content or known header-like labels)
            if (
                li.quantity == 0
                and li.unit_price == 0
                and li.total == 0
            ) or any(k in li.name for k in metadata_keywords):
                continue

            line_items.append(li)

        return line_items

    # ---- Raw text heuristic ----
    def _extract_line_items_from_text(self, text: str) -> List[LineItem]:
        line_items: List[LineItem] = []
        if not text:
            return line_items

        metadata_keywords = [
            "internal po reference",
            "po number",
            "customer po number",
            "linked customer po",
            "issue date",
            "po date",
            "reference",
        ]

        for raw in text.splitlines():
            line = raw.strip()
            if len(line) < 5:
                continue

            lower = line.lower()
            if any(k in lower for k in ["po no", "order date", "grand total", "subtotal"]):
                continue

            nums = re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", line)
            if not nums:
                continue

            # Heuristic: last number is total, second-last is unit, earlier may be qty
            total = self.num_norm.normalize_number(nums[-1])
            unit = self.num_norm.normalize_number(nums[-2]) if len(nums) >= 2 else 0.0
            qty = self.num_norm.normalize_number(nums[-3]) if len(nums) >= 3 else 0.0

            if qty == 0 and unit > 0 and total > 0:
                qty = total / unit

            # Extract non-numeric part as name
            name = re.sub(r"\d+(?:,\d{3})*(?:\.\d+)?", "", line)
            name = re.sub(r"Rs\.?|INR|USD|\$|€", "", name, flags=re.IGNORECASE)

            li = LineItem(
                name=self.text_norm.normalize_name(name),
                quantity=int(qty) if qty else 0,
                unit_price=unit,
                total=total if total else unit * qty,
            )
            # Skip metadata-only lines
            if not li.name:
                continue
            if (
                li.quantity == 0
                and li.unit_price == 0
                and li.total == 0
            ) or any(k in li.name for k in metadata_keywords):
                continue

            line_items.append(li)

        return line_items

    # ---- Totals (OTC / ARC / Grand Total) ----
    def _extract_totals(self, text: str, tables: List[List[str]]) -> Tuple[float, float, float]:
        otc = arc = grand = 0.0

        # Helper to check row/cell text for keywords
        def has_keyword(s: str, keywords: List[str]) -> bool:
            s = s.lower()
            return any(k in s for k in keywords)

        # 1) Inspect tables for rows that look like totals (preferred)
        if tables:
            for row in tables:
                joined = " ".join(str(c) for c in row).lower()
                # Skip rows with no numeric content
                if not re.search(r"\d", joined):
                    continue

                if has_keyword(joined, ["one time", "one-time", "otc"]):
                    val = self._extract_amount_from_line(joined)
                    if val:
                        otc = max(otc, val)
                if has_keyword(joined, ["annual", "recurring", "arc"]):
                    val = self._extract_amount_from_line(joined)
                    if val:
                        arc = max(arc, val)
                if has_keyword(joined, ["grand total"]) or ("total" in joined and any(k in joined for k in ["grand", "net payable", "net payable", "net due"])):
                    val = self._extract_amount_from_line(joined)
                    if val:
                        grand = max(grand, val)

        # 2) Fallback to scanning raw text lines
        if text:
            for line in text.splitlines():
                lower = line.lower()
                if has_keyword(lower, ["one time", "one-time", "otc"]):
                    val = self._extract_amount_from_line(line)
                    if val:
                        otc = max(otc, val)
                if has_keyword(lower, ["annual", "recurring", "arc"]):
                    val = self._extract_amount_from_line(line)
                    if val:
                        arc = max(arc, val)
                if "grand total" in lower or ("total" in lower and any(k in lower for k in ["grand", "net payable", "net due"])):
                    val = self._extract_amount_from_line(line)
                    if val:
                        grand = max(grand, val)

        return otc, arc, grand

    def _extract_amount_from_line(self, line: str) -> float:
        # Find candidate numeric tokens
        raw_nums = re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", line)
        if not raw_nums:
            return 0.0

        candidates: List[float] = []
        for tok in raw_nums:
            try:
                val = self.num_norm.normalize_number(tok)
            except Exception:
                continue
            # Filter out likely years (1900-2100) when token had no decimal point
            if re.fullmatch(r"\d{4}", tok) and 1900 <= int(tok) <= 2100:
                continue
            candidates.append(val)

        if not candidates:
            # As fallback, try normalizing the last token (legacy behavior)
            return self.num_norm.normalize_number(raw_nums[-1])

        # Prefer values with decimal portion (monetary) if present
        decimals = [v for v in candidates if abs(v - int(v)) > 0.0]
        if decimals:
            return max(decimals)

        # Otherwise pick the largest numeric candidate (totals usually the largest number on the row)
        return max(candidates)


# ---------------------- Comparator (Guardrails) ---------------------- #

class POComparator:
    """Compares two PurchaseOrder objects with 6-point audit + guardrails."""

    def __init__(self):
        self.text_norm = TextNormalizer()
        self.num_norm = NumericNormalizer()

    def compare(self, a: PurchaseOrder, b: PurchaseOrder) -> Dict[str, Any]:
        # 1. Product-level matching
        matches, product_overlap, product_denom = self._match_line_items(a.line_items, b.line_items)

        # 2. Adaptive gate (70% default). For very small orders lower the bar.
        threshold = 0.70
        if product_denom <= 2:
            threshold = 0.30
        if product_overlap < threshold:
            return {
                "status": "error",
                "reason": "Product overlap below threshold. Likely different orders.",
                "product_overlap": product_overlap,
                "threshold_used": threshold,
            }

        # 3. Per-line conflicts
        conflicts = []
        for la, lb, sim in matches:
            if la is None or lb is None:
                conflicts.append({
                    "type": "missing_item",
                    "company_item": la.name if la else None,
                    "customer_item": lb.name if lb else None,
                    "similarity": sim,
                })
                continue

            line_conf = self._compare_line_items(la, lb, sim)
            if line_conf:
                conflicts.append(line_conf)

        # 4. Totals comparison (strict: OTC vs OTC, ARC vs ARC, Grand vs Grand)
        total_conflicts = self._compare_totals(a, b)
        conflicts.extend(total_conflicts)

        return {
            "status": "conflict" if conflicts else "success",
            "product_overlap": product_overlap,
            "conflicts": conflicts,
        }

    # ---- Product matching ----
    def _match_line_items(
        self, items_a: List[LineItem], items_b: List[LineItem]
    ) -> Tuple[List[Tuple[Optional[LineItem], Optional[LineItem], float]], float, int]:
        matches: List[Tuple[Optional[LineItem], Optional[LineItem], float]] = []
        if not items_a or not items_b:
            return matches, 0.0

        names_a = [li.name for li in items_a]
        names_b = [li.name for li in items_b]

        used_b = set()
        for i, na in enumerate(names_a):
            # Find best match in B
            best_name, score, idx = process.extractOne(
                na, names_b, scorer=fuzz.token_set_ratio, score_cutoff=0
            )
            sim = score / 100.0
            # Only match if similarity is reasonable (>= 0.3) to avoid wrong matches
            # like "sify technologies" vs "ill connectivity"
            if idx is not None and idx not in used_b and sim >= 0.3:
                used_b.add(idx)
                matches.append((items_a[i], items_b[idx], sim))
            else:
                matches.append((items_a[i], None, 0.0))

        # Exclude common non-product keyword lines (OTC/ARC/Totals) from overlap calculation
        non_product_keys = {"otc", "arc", "tax", "vat", "total", "subtotal", "grand_total", "po value", "net payable"}
        def is_product_item(li: LineItem) -> bool:
            if not li or not li.name:
                return False
            name = li.name.lower()
            return not any(k in name for k in non_product_keys)

        prod_a_count = sum(1 for i in items_a if is_product_item(i))
        prod_b_count = sum(1 for i in items_b if is_product_item(i))
        denom = max(prod_a_count, prod_b_count) if max(prod_a_count, prod_b_count) > 0 else 1

        # Count only matches that pair product items with sufficient similarity
        # Use lower threshold (0.5) for product overlap to handle single-item orders
        good_matches = sum(1 for a, b, s in matches if b is not None and s >= 0.5 and is_product_item(a) and is_product_item(b))
        overlap = good_matches / denom if denom else 0.0

        return matches, overlap, denom

    # ---- Line item comparison ----
    def _compare_line_items(
        self, a: LineItem, b: LineItem, similarity: float
    ) -> Optional[Dict[str, Any]]:
        issues = []

        # If totals are equal (or extremely close), treat the line as financially
        # equivalent even if quantity/unit are expressed differently (e.g. qty 1 vs 0,
        # unit 500 vs 0 but same total). This matches the user's expectation that
        # same final charge with different formatting should not be a conflict.
        totals_close = abs(a.total - b.total) <= 0.01 and (a.total != 0 or b.total != 0)

        if not totals_close:
            if a.quantity != b.quantity:
                issues.append({
                    "field": "quantity",
                    "company": a.quantity,
                    "customer": b.quantity,
                })

            if abs(a.unit_price - b.unit_price) > 0.01:
                issues.append({
                    "field": "unit_price",
                    "company": a.unit_price,
                    "customer": b.unit_price,
                })

        # Always check total; if totals differ, it's a real financial conflict
        if abs(a.total - b.total) > 0.01:
            issues.append({
                "field": "total",
                "company": a.total,
                "customer": b.total,
            })

        if not issues:
            return None

        return {
            "type": "line_item_mismatch",
            "company_item": a.name,
            "customer_item": b.name,
            "similarity": similarity,
            "issues": issues,
        }

    # ---- Totals comparison ----
    def _compare_totals(self, a: PurchaseOrder, b: PurchaseOrder) -> List[Dict[str, Any]]:
        issues = []

        if abs(a.otc - b.otc) > 0.01:
            issues.append({
                "type": "otc_mismatch",
                "company": a.otc,
                "customer": b.otc,
            })

        if abs(a.arc - b.arc) > 0.01:
            issues.append({
                "type": "arc_mismatch",
                "company": a.arc,
                "customer": b.arc,
            })

        if abs(a.grand_total - b.grand_total) > 0.01:
            issues.append({
                "type": "grand_total_mismatch",
                "company": a.grand_total,
                "customer": b.grand_total,
            })

        return issues


# ---------------------- High-level API ---------------------- #

def audit_two_files(path_a: str, path_b: str) -> Dict[str, Any]:
    """
    End-to-end workflow:

    1. Extracts PO A and PO B from any supported format
    2. Normalizes into unified schema
    3. Runs comparison with 70% gate and strict OTC/ARC/Grand comparison
    4. Returns conflict report
    """
    extractor = DocumentExtractor()
    po_a = extractor.extract_po(path_a, order_id="COMPANY_PO")
    po_b = extractor.extract_po(path_b, order_id="CUSTOMER_PO")

    comparator = POComparator()
    result = comparator.compare(po_a, po_b)
    result["company_po"] = po_a
    result["customer_po"] = po_b
    return result


if __name__ == "__main__":
    # Example usage
    import json
    res = audit_two_files("Company_PO_Bookstore.pdf", "Customer_PO_Table_Amount_In_Words1.pdf")
    print(json.dumps(res, default=lambda o: o.__dict__, indent=2))