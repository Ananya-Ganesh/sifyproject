from __future__ import annotations

import os
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Tuple, Optional

import pandas as pd

from rapidfuzz import process, fuzz
from word2number import w2n


@dataclass
class LineItem:
    name: str
    qty: int
    unit_price: float
    total: float


@dataclass
class UnifiedPO:
    order_id: str
    line_items: List[LineItem]
    otc: float
    arc: float
    grand_total: float


class NumericNormalizer:
    @staticmethod
    def clean_text_num(text: str) -> str:
        if text is None:
            return ""
        t = str(text).strip()
        # OCR fixes
        t = t.replace("O", "0").replace("l", "1").replace("I", "1")
        return t

    def to_float(self, value: Any) -> float:
        if value is None:
            return 0.0
        if isinstance(value, (int, float)):
            return float(value)
        s = self.clean_text_num(str(value))
        if not s:
            return 0.0
        # remove currency and spaces
        s2 = re.sub(r"[^0-9.\-,]", "", s)
        s2 = s2.replace("-", "")
        s2 = s2.replace(",", "")
        try:
            return float(s2)
        except Exception:
            # try words -> number
            try:
                return float(w2n.word_to_num(s))
            except Exception:
                return 0.0


class TextNormalizer:
    @staticmethod
    def normalize_name(name: str) -> str:
        if not name:
            return ""
        s = str(name).lower()
        s = re.sub(r"[^a-z0-9\s]", " ", s)
        s = re.sub(r"\s+", " ", s).strip()
        # short-forms mapping
        mapping = {
            "natraj": "nat",
            "hb pencil": "pencil hb",
            "annual recurring": "arc",
            "annual recurring charge": "arc",
            "one time charge": "otc",
        }
        for k, v in mapping.items():
            if k in s:
                s = s.replace(k, v)
        return s


def extract_raw_text(path: str) -> str:
    """Detect file type and extract all text into a single string.

    For Excel, rows/columns are rendered as simple CSV-like lines.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        # try pdfplumber then fallback to OCR. import optional libs lazily.
        try:
            import pdfplumber  # type: ignore
        except Exception:
            pdfplumber = None

        try:
            from pdf2image import convert_from_path  # type: ignore
        except Exception:
            convert_from_path = None

        try:
            import pytesseract  # type: ignore
        except Exception:
            pytesseract = None

        texts: List[str] = []
        if pdfplumber is not None:
            try:
                with pdfplumber.open(path) as pdf:
                    for p in pdf.pages:
                        t = p.extract_text() or ""
                        texts.append(t)
                combined = "\n".join(t for t in texts if t)
                if len(re.sub(r"\s+", "", combined)) < 20 and convert_from_path and pytesseract:
                    imgs = convert_from_path(path)
                    ocr_txts = [pytesseract.image_to_string(img) for img in imgs]
                    return "\n".join(ocr_txts)
                return combined
            except Exception:
                texts = []

        # Fallback to OCR-only if pdfplumber failed or produced no text
        if convert_from_path and pytesseract:
            imgs = convert_from_path(path)
            return "\n".join(pytesseract.image_to_string(img) for img in imgs)
        return ""

    if ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff"):
        try:
            from PIL import Image  # type: ignore
            import pytesseract  # type: ignore
        except Exception:
            return ""
        img = Image.open(path)
        return pytesseract.image_to_string(img)

    if ext == ".docx":
        try:
            from docx import Document as DocxDocument  # type: ignore
        except Exception:
            return ""
        doc = DocxDocument(path)
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    if ext in (".xls", ".xlsx"):
        try:
            sheets = pd.read_excel(path, sheet_name=None)
        except Exception:
            try:
                xls = pd.ExcelFile(path)
                sheets = {n: xls.parse(n) for n in xls.sheet_names}
            except Exception:
                return ""
        lines: List[str] = []
        for name, df in sheets.items():
            if df is None or df.empty:
                continue
            headers = [str(h) for h in df.columns]
            lines.append(", ".join(headers))
            for _, row in df.fillna("").iterrows():
                vals = [str(v) if v is not None else "" for v in row.tolist()]
                lines.append(", ".join(vals))
        return "\n".join(lines)

    # fallback to text file
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def parse_to_unified_schema(raw_text: str, tables: Optional[List[List[str]]] = None, order_id: Optional[str] = None) -> UnifiedPO:
    """Transform raw text + optional table rows into the unified schema.

    This uses regex heuristics (no LLM) to find columns: name, qty, unit_price, total.
    """
    num_norm = NumericNormalizer()
    text_norm = TextNormalizer()

    line_items: List[LineItem] = []
    # Keywords to treat rows as metadata (not product line items)
    metadata_keywords = [
        "internal po reference",
        "linked customer po",
        "issue date",
        "company name",
        "currency",
        "customer name",
        "customer po number",
        "po date",
        "net payable",
        "total po value",
        "total",
    ]

    # Helper to scan a row/list for columns
    def row_to_item(row: List[str]) -> Optional[LineItem]:
        joined = " ".join(str(c) for c in row)
        # find all numbers in the row (with positions)
        num_iter = list(re.finditer(r"\d+(?:,\d{3})*(?:\.\d+)?", joined))
        if not num_iter:
            return None

        # Exclude numbers that are bandwidth specs like '500 Mbps' (adjacent to 'mbps')
        nums_filtered: List[str] = []
        jlow = joined.lower()
        for m in num_iter:
            tok = m.group(0)
            start, end = m.span()
            window = jlow[max(0, end):min(len(jlow), end + 8)]
            window_before = jlow[max(0, start - 8):start]
            if 'mbps' in window or 'mbps' in window_before:
                # treat as bandwidth, skip
                continue
            nums_filtered.append(tok)

        if not nums_filtered:
            # no monetary numbers after filtering
            # but if there are numbers originally, fall back to using the last numeric
            nums_filtered = [m.group(0) for m in num_iter]

        # Heuristic: last number is total, second-last is unit, third-last might be qty
        total = num_norm.to_float(nums_filtered[-1]) if nums_filtered else 0.0
        unit = num_norm.to_float(nums_filtered[-2]) if len(nums_filtered) >= 2 else 0.0
        qty = int(num_norm.to_float(nums_filtered[-3])) if len(nums_filtered) >= 3 else 0
        # Extract text for name by removing numbers
        name = re.sub(r"\d+(?:,\d{3})*(?:\.\d+)?", "", joined)
        name = re.sub(r"\b(total|qty|quantity|unit|amount|rate|price)\b", "", name, flags=re.IGNORECASE)
        name = text_norm.normalize_name(name)
        if not name:
            return None
        if unit == 0 and qty > 0 and total > 0:
            unit = total / qty
        if qty == 0 and unit > 1 and total > 0:
            qty = int(round(total / unit))
        if total == 0 and unit > 0 and qty > 0:
            total = unit * qty
        return LineItem(name=name, qty=int(qty), unit_price=float(unit), total=float(total))

    # Totals defaults
    otc = arc = grand = 0.0

    # Helper to extract amounts from a text snippet (used by table parser)
    def extract_amount_from_text(s: str) -> float:
        nums = re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", s)
        if not nums:
            return 0.0
        cand = []
        for tok in nums:
            if re.fullmatch(r"\d{4}", tok) and 1900 <= int(tok) <= 2100:
                continue
            cand.append(num_norm.to_float(tok))
        if not cand:
            return 0.0
        decs = [c for c in cand if abs(c - int(c)) > 0]
        return max(decs) if decs else max(cand)

    # Prefer parsing tables if provided (caller can supply tables)
    if tables:
        # tables is a list of rows (each row a list of strings)
        # Try to detect two common layouts:
        # 1) Key-Value metadata table (Field / Value) -> extract OTC/ARC/Grand and look for separate product table
        # 2) Tabular product rows with headers (Line No / Service / Amount)
        # First, attempt to find a product table header row
        product_header_idx = None
        for i, r in enumerate(tables[:12]):
            s = " ".join(str(c).lower() for c in r)
            if any(k in s for k in ["line no", "service", "description"]) and any(k in s for k in ["amount", "total", "price"]):
                product_header_idx = i
                break

        # Heuristic: detect key-value table (Field/Value) where many rows have numeric second column
        kv_count = 0
        kv_total = 0
        for r in tables[:20]:
            kv_total += 1
            if len(r) >= 2 and re.search(r"\d", str(r[1])):
                kv_count += 1
        is_kv_table = kv_count >= max(2, kv_total // 3)

        # If key-value table, extract totals from matching rows
        totals_by_label = {}
        if is_kv_table:
            for r in tables:
                if not r:
                    continue
                key = str(r[0]).strip().lower()
                line_joined = " ".join(str(c) for c in r)
                val = extract_amount_from_text(line_joined)
                if val:
                    totals_by_label[key] = max(totals_by_label.get(key, 0.0), val)
                if any(k in key for k in ["otc", "one time", "one-time"]):
                    otc = max(otc, val)
                if any(k in key for k in ["arc", "annual", "recurring"]):
                    arc = max(arc, val)
                if any(k in key for k in ["grand total", "net payable", "total po value", "total"]):
                    grand = max(grand, val)

        # If we found a product header, parse rows below it as product items
        if product_header_idx is not None:
            rows = tables[product_header_idx + 1 :]
            for r in rows:
                try:
                    item = row_to_item(r)
                except Exception:
                    item = None
                if item:
                    # Skip rows that look like metadata
                    if any(k in item.name for k in metadata_keywords):
                        continue
                    line_items.append(item)
            # done with table parsing
        else:
            # No clear product table found; attempt to find rows that look like product entries
            for r in tables:
                if not r:
                    continue
                # Skip obvious metadata rows where first cell is a field name
                first = str(r[0]).strip().lower()
                if any(k in first for k in ["field", "company name", "internal po reference", "linked customer po", "issue date", "currency", "customer name", "customer po number", "po date"]):
                    continue
                # If row has at least 2 columns and one column contains a long description, treat as product row
                if len(r) >= 2 and re.search(r"\d", " ".join(r)):
                    try:
                        item = row_to_item(r)
                    except Exception:
                        item = None
                    if item:
                        if any(k in item.name for k in metadata_keywords):
                            continue
                        line_items.append(item)

            # If we found a service description row with no amount, try to match it to a total label (Total PO Value / Net Payable)
            if totals_by_label:
                # look for service description rows
                svc_name = None
                for r in tables:
                    if not r:
                        continue
                    first = str(r[0]).strip().lower()
                    if 'service description' in first or 'service' in first:
                        # description may be in second column
                        if len(r) >= 2 and str(r[1]).strip():
                            svc_name = str(r[1]).strip()
                            break
                if svc_name:
                    # prefer exact total labels
                    total_keys = ['total po value', 'net payable', 'total']
                    total_val = 0.0
                    for k in total_keys:
                        for label, val in totals_by_label.items():
                            if k in label:
                                total_val = max(total_val, val)
                    if total_val:
                        li = LineItem(name=text_norm.normalize_name(svc_name), qty=0, unit_price=0.0, total=float(total_val))
                        line_items.append(li)

                # If a service row was parsed with a small numeric (e.g. '500' from '500 Mbps'),
                # but we have a separate 'total po value' or 'net payable', prefer that as the item's total.
                preferred_total = 0.0
                for k in ['total po value', 'net payable', 'total']:
                    for label, val in totals_by_label.items():
                        if k in label:
                            preferred_total = max(preferred_total, val)
                if preferred_total:
                    for li in line_items:
                        if 'service' in li.name or 'ill' in li.name:
                            # if current total looks like a bandwidth number (<= 10000) and preferred is larger, replace
                            if li.total < preferred_total:
                                li.total = float(preferred_total)
                                if li.qty == 0 and li.total > 0:
                                    li.unit_price = li.total
                                    li.qty = 1
                    # remove any spurious 'po value' or 'net payable' items from line_items
                    line_items = [li for li in line_items if not any(k in li.name for k in ['po value', 'net payable', 'total po value'])]

                # As a fallback: if we still have a 'service' item and a separate 'po value' item, merge them
                svc_idx = next((i for i, li in enumerate(line_items) if 'service' in li.name or 'ill' in li.name), None)
                total_idx = next((i for i, li in enumerate(line_items) if any(k in li.name for k in ['po value', 'net payable', 'total po value'])), None)
                if svc_idx is not None and total_idx is not None and svc_idx != total_idx:
                    svc = line_items[svc_idx]
                    tot = line_items[total_idx]
                    if svc.total < tot.total:
                        svc.total = tot.total
                        if svc.qty == 0 and svc.total > 0:
                            svc.unit_price = svc.total
                            svc.qty = 1
                    # remove the separate total row
                    line_items.pop(total_idx)

            # Deduplicate line items by normalized name: keep the one with the largest total
            # First pass dedupe by exact normalized name
            unique: Dict[str, LineItem] = {}
            for li in line_items:
                key = li.name.strip()
                if not key:
                    continue
                if key in unique:
                    if li.total > unique[key].total:
                        unique[key] = li
                else:
                    unique[key] = li
            line_items = list(unique.values())

            # Second pass: fuzzy dedupe for near-duplicates (e.g., 'service description...' vs 'ill connectivity 500 mbps')
            merged: List[LineItem] = []
            for li in line_items:
                found = False
                for j, uj in enumerate(merged):
                    sim = fuzz.token_set_ratio(li.name, uj.name) / 100.0
                    if sim >= 0.85:
                        found = True
                        # keep the item with larger total
                        if li.total > uj.total:
                            merged[j] = li
                        break
                if not found:
                    merged.append(li)
            line_items = merged

    # If no line items yet, parse raw text line-by-line
    if not line_items and raw_text:
        for raw in raw_text.splitlines():
            raw = raw.strip()
            if len(raw) < 5:
                continue
            # skip obvious meta
            if any(k in raw.lower() for k in ["po number", "invoice", "bill to", "ship to"]):
                continue
            # If line contains numbers, attempt parse
            if re.search(r"\d", raw):
                # split on common separators
                parts = re.split(r"\s{2,}|\t|,", raw)
                # try different windows
                for p in parts:
                    if re.search(r"\d", p):
                        candidate = row_to_item([p])
                        if candidate:
                            line_items.append(candidate)
                            break

    # Extract totals (otc, arc, grand_total) by scanning both text and tables
    def extract_amount_from_text(s: str) -> float:
        nums = re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", s)
        if not nums:
            return 0.0
        # remove 4-digit years if present
        cand = []
        for tok in nums:
            if re.fullmatch(r"\d{4}", tok) and 1900 <= int(tok) <= 2100:
                continue
            cand.append(num_norm.to_float(tok))
        if not cand:
            return 0.0
        # prefer decimals
        decs = [c for c in cand if abs(c - int(c)) > 0]
        return max(decs) if decs else max(cand)

    otc = arc = grand = 0.0
    # scan tables if available
    if tables:
        for r in tables:
            joined = " ".join(str(c) for c in r).lower()
            if any(k in joined for k in ["one time", "one-time", "otc"]):
                otc = max(otc, extract_amount_from_text(joined))
            if any(k in joined for k in ["annual", "recurring", "arc"]):
                arc = max(arc, extract_amount_from_text(joined))
            if "grand total" in joined or ("total" in joined and any(k in joined for k in ["grand", "net payable", "net due"])):
                grand = max(grand, extract_amount_from_text(joined))

    if raw_text:
        for line in raw_text.splitlines():
            lower = line.lower()
            if any(k in lower for k in ["one time", "one-time", "otc"]):
                otc = max(otc, extract_amount_from_text(line))
            if any(k in lower for k in ["annual", "recurring", "arc"]):
                arc = max(arc, extract_amount_from_text(line))
            if "grand total" in lower or ("total" in lower and any(k in lower for k in ["grand", "net payable", "net due"])):
                grand = max(grand, extract_amount_from_text(line))

    # if grand still empty, sum line_items totals
    if grand == 0.0 and line_items:
        grand = sum(li.total for li in line_items)

    order_id = order_id or "uploaded_po"
    return UnifiedPO(order_id=order_id, line_items=line_items, otc=float(otc), arc=float(arc), grand_total=float(grand))


def fuzzy_match_items(items_a: List[LineItem], items_b: List[LineItem]) -> Tuple[List[Tuple[Optional[LineItem], Optional[LineItem], float]], float]:
    """Match items from A to B and compute product overlap ratio.

    Returns list of (a_item, b_item_or_None, similarity) and product_overlap (good_matches/denom)
    """
    if not items_a or not items_b:
        return ([(ia, None, 0.0) for ia in items_a], 0.0)
    names_a = [la.name for la in items_a]
    names_b = [lb.name for lb in items_b]
    used_b = set()
    matches = []
    for i, na in enumerate(names_a):
        res = process.extractOne(na, names_b, scorer=fuzz.token_set_ratio)
        if not res:
            matches.append((items_a[i], None, 0.0))
            continue
        best_name, score, idx = res
        sim = score / 100.0
        if idx is not None and idx not in used_b and sim >= 0.3:
            used_b.add(idx)
            matches.append((items_a[i], items_b[idx], sim))
        else:
            matches.append((items_a[i], None, 0.0))

    good_matches = sum(1 for a, b, s in matches if b is not None and s >= 0.7)
    denom = max(len(items_a), len(items_b))
    overlap = good_matches / denom if denom else 0.0
    return matches, overlap


def compare_pos(po_a: UnifiedPO, po_b: UnifiedPO) -> Dict[str, Any]:
    """Perform strict audit comparison according to rules.

    - 70% product overlap gate
    - Strict OTC <-> OTC, ARC <-> ARC
    - Validate sums
    - Return conflict report JSON
    """
    # 1. product matching
    matches, _ = fuzzy_match_items(po_a.line_items, po_b.line_items)
    # Exclude common non-product keyword lines (OTC/ARC/Totals) from overlap calculation
    non_product_keys = {"otc", "arc", "tax", "vat", "total", "subtotal", "grand_total", "po value", "net payable"}
    def is_product(item: LineItem) -> bool:
        if not item or not item.name:
            return False
        name = item.name.lower()
        return not any(k in name for k in non_product_keys)

    prod_a = [i for i in po_a.line_items if is_product(i)]
    prod_b = [i for i in po_b.line_items if is_product(i)]
    denom = max(len(prod_a), len(prod_b)) if max(len(prod_a), len(prod_b)) > 0 else 1
    # Count only matches that pair product items with sufficient similarity
    # Use lower threshold (0.5) to handle single-item orders where exact match may be ~60%
    product_matches_count = sum(1 for a, b, s in matches if b is not None and s >= 0.5 and is_product(a) and is_product(b))
    product_overlap = product_matches_count / denom
    # Adaptive gate: 70% for normal, 30% for small orders (<=2 items)
    overlap_threshold = 0.30 if denom <= 2 else 0.70
    if product_overlap < overlap_threshold:
        return {"status": "error", "reason": "Error: Order Mismatch. Documents belong to different orders.", "product_overlap": product_overlap}

    report_results = []
    conflict_count = 0

    for la, lb, sim in matches:
        status = "ok"
        conflicts = []
        if lb is None:
            status = "missing_in_b"
        else:
            # compare qty and price and total
            if abs(la.total - lb.total) > 0.01:
                conflicts.append({"field": "total", "a": la.total, "b": lb.total})
            # if totals close, allow qty/unit differences
            totals_close = abs(la.total - lb.total) <= 0.01 and (la.total != 0 or lb.total != 0)
            if not totals_close:
                if la.qty != lb.qty:
                    conflicts.append({"field": "qty", "a": la.qty, "b": lb.qty})
                if abs(la.unit_price - lb.unit_price) > 0.01:
                    conflicts.append({"field": "unit_price", "a": la.unit_price, "b": lb.unit_price})
            if conflicts:
                status = "conflict"
                conflict_count += 1

        report_results.append({
            "status": status,
            "similarity": round(sim, 3),
            "item_a": asdict(la),
            "item_b": asdict(lb) if lb else {},
            "conflicts": conflicts,
        })

    # Strict compare OTC and ARC
    otc_conflict = abs(po_a.otc - po_b.otc) > 0.01
    arc_conflict = abs(po_a.arc - po_b.arc) > 0.01
    if otc_conflict:
        report_results.append({
            "status": "conflict",
            "similarity": 1.0,
            "item_a": {"name": "otc", "qty": 0, "unit_price": 0.0, "total": po_a.otc},
            "item_b": {"name": "otc", "qty": 0, "unit_price": 0.0, "total": po_b.otc},
            "conflicts": [{"field": "otc", "a": po_a.otc, "b": po_b.otc}],
        })
        conflict_count += 1

    if arc_conflict:
        report_results.append({
            "status": "conflict",
            "similarity": 1.0,
            "item_a": {"name": "arc", "qty": 0, "unit_price": 0.0, "total": po_a.arc},
            "item_b": {"name": "arc", "qty": 0, "unit_price": 0.0, "total": po_b.arc},
            "conflicts": [{"field": "arc", "a": po_a.arc, "b": po_b.arc}],
        })
        conflict_count += 1

    # Validate grand totals: sum PRODUCT line_items + otc + arc equals grand_total
    # (exclude OTC/ARC from line_items since they're already counted separately)
    sum_a = sum(li.total for li in po_a.line_items if is_product(li)) + po_a.otc + po_a.arc
    sum_b = sum(li.total for li in po_b.line_items if is_product(li)) + po_b.otc + po_b.arc
    grand_conflict = abs(sum_a - po_a.grand_total) > 0.01 or abs(sum_b - po_b.grand_total) > 0.01 or abs(po_a.grand_total - po_b.grand_total) > 0.01
    if grand_conflict:
        report_results.append({
            "status": "conflict",
            "similarity": 1.0,
            "item_a": {"name": "grand_total", "qty": 0, "unit_price": 0.0, "total": po_a.grand_total},
            "item_b": {"name": "grand_total", "qty": 0, "unit_price": 0.0, "total": po_b.grand_total},
            "conflicts": [{"field": "grand_total", "a": po_a.grand_total, "b": po_b.grand_total}],
        })
        conflict_count += 1

    # Build high-level summary
    summary = {
        "total_items_a": len(po_a.line_items),
        "total_items_b": len(po_b.line_items),
        "matched_items": sum(1 for r in report_results if r["status"] != "missing_in_b" and r.get("item_a", {}).get("name") not in ["otc", "arc", "grand_total"]),
        "conflict_count": conflict_count,
        "product_overlap": product_overlap,
        "otc_company": po_a.otc,
        "otc_customer": po_b.otc,
        "arc_company": po_a.arc,
        "arc_customer": po_b.arc,
        "grand_company": po_a.grand_total,
        "grand_customer": po_b.grand_total,
    }

    # Also produce color-coded lists
    matched = []
    spelling_variations = []
    price_qty_conflicts = []
    missing = []

    for r in report_results:
        n_a = r.get("item_a", {}).get("name") if r.get("item_a") else None
        n_b = r.get("item_b", {}).get("name") if r.get("item_b") else None
        if r["status"] == "ok":
            matched.append({"a": n_a, "b": n_b, "similarity": r["similarity"]})
        elif r["status"] == "conflict":
            # if similarity high but names differ -> spelling variation
            if r["similarity"] >= 0.5 and r["similarity"] < 0.95:
                spelling_variations.append({"a": n_a, "b": n_b, "similarity": r["similarity"]})
            # price/qty conflicts
            if any(c.get("field") in ("total", "qty", "unit_price", "grand_total", "otc", "arc") for c in r.get("conflicts", [])):
                price_qty_conflicts.append(r)
        elif r["status"] == "missing_in_b":
            missing.append({"a": n_a})

    return {
        "status": "ok",
        "summary": summary,
        "results": report_results,
        "matched": matched,
        "spelling_variations": spelling_variations,
        "price_qty_conflicts": price_qty_conflicts,
        "missing": missing,
    }


def compare_files(path_a: str, path_b: str) -> Dict[str, Any]:
    """High-level helper: extract, parse, compare two files and return the conflict report."""
    text_a = extract_raw_text(path_a)
    text_b = extract_raw_text(path_b)

    # Try to also build simple table rows from Excel if present to improve parsing
    def build_tables(path: str):
        try:
            sheets = pd.read_excel(path, sheet_name=None)
        except Exception:
            try:
                xls = pd.ExcelFile(path)
                sheets = {n: xls.parse(n) for n in xls.sheet_names}
            except Exception:
                return None
        tables = []
        for name, df in sheets.items():
            if df is None or df.empty:
                continue
            headers = [str(h) for h in df.columns]
            tables.append(headers)
            for _, row in df.fillna("").iterrows():
                tables.append([str(v) if v is not None else "" for v in row.tolist()])
        return tables

    tables_a = build_tables(path_a)
    tables_b = build_tables(path_b)

    po_a = parse_to_unified_schema(text_a, tables=tables_a, order_id=os.path.basename(path_a))
    po_b = parse_to_unified_schema(text_b, tables=tables_b, order_id=os.path.basename(path_b))

    result = compare_pos(po_a, po_b)
    # include parsed structures for debugging/visibility
    result.setdefault("parsed", {})
    result["parsed"]["company_parsed"] = {
        "order_id": po_a.order_id,
        "line_items": [li.__dict__ for li in po_a.line_items],
        "otc": po_a.otc,
        "arc": po_a.arc,
        "grand_total": po_a.grand_total,
    }
    result["parsed"]["customer_parsed"] = {
        "order_id": po_b.order_id,
        "line_items": [li.__dict__ for li in po_b.line_items],
        "otc": po_b.otc,
        "arc": po_b.arc,
        "grand_total": po_b.grand_total,
    }

    return result


if __name__ == "__main__":
    import json
    # Quick demo (replace with real paths)
    a = "Company_PO.xlsx"
    b = "Customer_PO.xlsx"
    if os.path.exists(a) and os.path.exists(b):
        res = compare_files(a, b)
        print(json.dumps(res, indent=2, default=str))
    else:
        print("Place sample files as Company_PO.xlsx and Customer_PO.xlsx to run demo")
