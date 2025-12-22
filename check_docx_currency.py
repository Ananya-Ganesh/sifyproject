#!/usr/bin/env python3
"""Check currency symbols in DOCX files"""
from docx import Document

file1 = r'c:\Users\Administrator\Downloads\Company_PO_BookstoreCH.docx'
file2 = r'c:\Users\Administrator\Downloads\Customer_Purchase_OrderCH.docx'

print("=" * 80)
print("FILE 1: Company PO")
print("=" * 80)
try:
    doc1 = Document(file1)
    text1 = '\n'.join([p.text for p in doc1.paragraphs])
    print(f"Text length: {len(text1)}")
    print(f"First 500 chars:\n{text1[:500]}")
    print()
    
    has_dollar = '$' in text1
    has_rupee = '\u20b9' in text1  # ₹ symbol
    print(f"Contains $: {has_dollar}")
    print(f"Contains rupee (₹): {has_rupee}")
    
    # Show all unique characters that look like currency
    special_chars = set()
    for char in text1:
        if ord(char) > 127:  # Non-ASCII
            special_chars.add(f"{char} (U+{ord(char):04X})")
    if special_chars:
        print(f"Special characters found: {special_chars}")
except Exception as e:
    print(f"Error reading file1: {e}")

print()
print("=" * 80)
print("FILE 2: Customer PO")
print("=" * 80)
try:
    doc2 = Document(file2)
    text2 = '\n'.join([p.text for p in doc2.paragraphs])
    print(f"Text length: {len(text2)}")
    print(f"First 500 chars:\n{text2[:500]}")
    print()
    
    has_dollar = '$' in text2
    has_rupee = '\u20b9' in text2  # ₹ symbol
    print(f"Contains $: {has_dollar}")
    print(f"Contains rupee (₹): {has_rupee}")
    
    # Show all unique characters that look like currency
    special_chars = set()
    for char in text2:
        if ord(char) > 127:  # Non-ASCII
            special_chars.add(f"{char} (U+{ord(char):04X})")
    if special_chars:
        print(f"Special characters found: {special_chars}")
except Exception as e:
    print(f"Error reading file2: {e}")
